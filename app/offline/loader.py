"""
文档加载器

支持从 PDF、DOCX、TXT 文件中加载药品说明书原文。

使用方式:
    from app.offline.loader import load_document, LoadedDocument

    doc = load_document("data/raw/阿司匹林说明书.pdf")
    print(doc.raw_text[:200])
    print(doc.inferred_drug_name)
"""

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from loguru import logger


# ============================================================
# 数据类
# ============================================================
@dataclass
class LoadedDocument:
    """加载后的文档容器"""

    raw_text: str
    source_file: str
    inferred_drug_name: Optional[str] = None
    file_type: Optional[str] = None  # "pdf" | "docx" | "txt"
    metadata: dict = field(default_factory=dict)


# ============================================================
# 各格式加载器
# ============================================================
def load_pdf(file_path: Path) -> str:
    """从 PDF 文件中提取文本（使用 pypdf）"""
    from pypdf import PdfReader

    logger.info(f"正在读取 PDF: {file_path.name}")
    reader = PdfReader(str(file_path))
    pages_text: list[str] = []

    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text()
        if text:
            # 清理页面内的多余空白
            text = text.strip()
            pages_text.append(text)
        else:
            logger.debug(f"  第 {i} 页无文本内容")

    full_text = "\n\n".join(pages_text)
    logger.info(f"PDF 读取完成: {len(reader.pages)} 页, {len(full_text)} 字符")
    return full_text


def load_docx(file_path: Path) -> str:
    """从 DOCX 文件中提取文本（使用 python-docx）"""
    from docx import Document as DocxDocument

    logger.info(f"正在读取 DOCX: {file_path.name}")
    doc = DocxDocument(str(file_path))
    paragraphs: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    full_text = "\n".join(paragraphs)
    logger.info(f"DOCX 读取完成: {len(doc.paragraphs)} 段落, {len(full_text)} 字符")
    return full_text


def load_txt(file_path: Path) -> str:
    """从 TXT 文件中读取文本（UTF-8 优先，回退 GBK）"""
    logger.info(f"正在读取 TXT: {file_path.name}")

    # 尝试 UTF-8
    try:
        text = file_path.read_text(encoding="utf-8")
        logger.info(f"TXT (UTF-8) 读取完成: {len(text)} 字符")
        return text
    except UnicodeDecodeError:
        logger.debug("UTF-8 解码失败，尝试 GBK...")

    # 回退 GBK（中文 Windows 常见编码）
    try:
        text = file_path.read_text(encoding="gbk", errors="replace")
        logger.info(f"TXT (GBK) 读取完成: {len(text)} 字符")
        return text
    except Exception as e:
        raise LoaderError(f"TXT 文件编码无法识别: {file_path}") from e


# ============================================================
# 药名推断
# ============================================================
# 常见的文件名后缀（需要去除以提取药名）
_DRUG_NAME_STRIP_PATTERNS = [
    r"说明书.*$",  # "说明书.pdf", "说明书_拜耳.txt"
    r"药品?说?明书.*$",
    r"用药指南.*$",
    r"\(.*?\)",  # "(拜耳医药)" 等括号内容，保留药名
    r"（.*?）",  # 中文括号
    r"_\d{4}.*$",  # "_2024版" 等年份后缀
    r"[-_][一-鿿]*公司.*$",  # "-拜耳医药公司"
    r"\.(pdf|docx?|txt)$",  # 扩展名（通常不存在，因为已经 split）
]


def infer_drug_name(file_path: Path) -> Optional[str]:
    """
    从文件名推断药品名称。

    Args:
        file_path: 文档文件路径

    Returns:
        推断的药品名称，如果无法推断则返回 None。

    Examples:
        "阿司匹林肠溶片说明书_拜耳医药.pdf" -> "阿司匹林肠溶片"
        "布洛芬缓释胶囊.doc" -> "布洛芬缓释胶囊"
    """
    # 去掉扩展名
    stem = file_path.stem

    # 逐个应用清理规则
    name = stem
    for pattern in _DRUG_NAME_STRIP_PATTERNS:
        name = re.sub(pattern, "", name)

    name = name.strip().strip("_- （）()")

    if not name:
        logger.warning(f"无法从文件名推断药名: {file_path.name}")
        return None

    logger.info(f"从文件名推断药名: '{stem}' → '{name}'")
    return name


# ============================================================
# 异常
# ============================================================
class LoaderError(Exception):
    """文档加载错误"""

    pass


# ============================================================
# 公共 API
# ============================================================
# 支持的文件扩展名映射
_LOADER_MAP = {
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".doc": load_docx,  # .doc 也尝试用 python-docx 打开
    ".txt": load_txt,
}


def load_document(file_path: Path) -> LoadedDocument:
    """
    加载单个文档，根据扩展名自动选择加载器。

    Args:
        file_path: 文档路径（必须是 .pdf / .docx / .doc / .txt）

    Returns:
        LoadedDocument 对象（含原文、文件名、推断药名、文件类型）

    Raises:
        FileNotFoundError: 文件不存在
        LoaderError: 不支持的文件格式或读取失败
    """
    file_path = Path(file_path)

    if not file_path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")

    suffix = file_path.suffix.lower()
    if suffix not in _LOADER_MAP:
        raise LoaderError(
            f"不支持的文件格式: {suffix}。"
            f"支持的格式: {', '.join(_LOADER_MAP.keys())}"
        )

    loader = _LOADER_MAP[suffix]
    file_type = suffix.lstrip(".")

    try:
        raw_text = loader(file_path)
    except Exception as e:
        raise LoaderError(f"读取文件失败 ({file_path.name}): {e}") from e

    if not raw_text.strip():
        logger.warning(f"文件内容为空: {file_path.name}")

    # 推断药名
    drug_name = infer_drug_name(file_path)

    return LoadedDocument(
        raw_text=raw_text,
        source_file=str(file_path),
        inferred_drug_name=drug_name,
        file_type=file_type,
    )


def load_documents_from_dir(
    dir_path: Path,
    recursive: bool = False,
) -> list[LoadedDocument]:
    """
    加载目录中所有支持的文档文件。

    Args:
        dir_path: 目录路径
        recursive: 是否递归搜索子目录

    Returns:
        LoadedDocument 列表（跳过加载失败的文件）
    """
    dir_path = Path(dir_path)
    if not dir_path.is_dir():
        raise NotADirectoryError(f"目录不存在: {dir_path}")

    pattern = "**/*" if recursive else "*"
    supported_extensions = tuple(_LOADER_MAP.keys())

    documents: list[LoadedDocument] = []
    for file_path in sorted(dir_path.glob(pattern)):
        if file_path.suffix.lower() in supported_extensions:
            try:
                doc = load_document(file_path)
                documents.append(doc)
            except Exception as e:
                logger.error(f"跳过文件 {file_path.name}: {e}")

    logger.info(
        f"目录加载完成: {len(documents)} 个文档 (来自 {dir_path}, recursive={recursive})"
    )
    return documents
